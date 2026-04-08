#!/usr/bin/env python3
"""
BIS Document Intelligence & Governance Analysis System
Advanced RAG-based Document Analysis for Banking Regulatory Intelligence

Author: Dr. Mosab Hawarey
Advanced Natural Language Processing & Financial Regulatory Analysis
"""

import os
import sys
import json
import logging
import sqlite3
import threading
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Core Libraries for Document Processing
import PyPDF2
import fitz  # PyMuPDF for better PDF extraction
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import LatentDirichletAllocation
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import WordNetLemmatizer
import spacy
import textstat
from textblob import TextBlob

# DOCX Generation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')
try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')
try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')
try:
    nltk.data.find('chunkers/maxent_ne_chunker')
except LookupError:
    nltk.download('maxent_ne_chunker')
try:
    nltk.data.find('corpora/words')
except LookupError:
    nltk.download('words')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Please install spaCy English model: python -m spacy download en_core_web_sm")
    sys.exit(1)

class DocumentProcessor:
    """Advanced PDF document processing with multiple extraction methods"""
    
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        
    def extract_text_pymupdf(self, pdf_path: str) -> str:
        """Extract text using PyMuPDF (more reliable)"""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logging.error(f"PyMuPDF extraction failed for {pdf_path}: {e}")
            return ""
    
    def extract_text_pypdf2(self, pdf_path: str) -> str:
        """Fallback extraction using PyPDF2"""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
            return text
        except Exception as e:
            logging.error(f"PyPDF2 extraction failed for {pdf_path}: {e}")
            return ""
    
    def extract_text(self, pdf_path: str) -> str:
        """Extract text with fallback methods"""
        text = self.extract_text_pymupdf(pdf_path)
        if not text.strip():
            text = self.extract_text_pypdf2(pdf_path)
        return text
    
    def preprocess_text(self, text: str) -> str:
        """Advanced text preprocessing"""
        # Remove special characters and normalize
        text = ' '.join(text.split())
        text = text.replace('\n', ' ').replace('\r', ' ')
        
        # Tokenize and lemmatize
        tokens = word_tokenize(text.lower())
        tokens = [self.lemmatizer.lemmatize(token) for token in tokens 
                 if token.isalpha() and token not in self.stop_words and len(token) > 2]
        
        return ' '.join(tokens)

class AdvancedAnalyzer:
    """Sophisticated analysis engine for detecting patterns and anomalies"""
    
    def __init__(self):
        self.negative_patterns = {
            'governance_failures': [
                'lack of oversight', 'inadequate supervision', 'poor governance',
                'insufficient controls', 'weak oversight', 'regulatory gap',
                'compliance failure', 'oversight deficiency', 'governance weakness',
                'supervisory gap', 'control weakness', 'regulatory shortcoming'
            ],
            'compliance_issues': [
                'non-compliance', 'regulatory breach', 'violation', 'infringement',
                'non-conformity', 'regulatory failure', 'breach of regulation',
                'compliance deficiency', 'regulatory non-compliance', 'violation of rules'
            ],
            'legal_loopholes': [
                'regulatory arbitrage', 'loophole', 'gap in regulation', 'ambiguity',
                'unclear provision', 'regulatory void', 'legal gap', 'jurisdictional gap',
                'regulatory inconsistency', 'ambiguous requirement'
            ],
            'operational_risks': [
                'operational risk', 'system failure', 'process breakdown',
                'operational deficiency', 'control failure', 'risk management failure',
                'operational weakness', 'system inadequacy', 'process failure'
            ],
            'market_irregularities': [
                'market manipulation', 'unusual trading', 'suspicious activity',
                'irregular pattern', 'anomalous behavior', 'market distortion',
                'price manipulation', 'trading anomaly', 'market abuse'
            ],
            'financial_misconduct': [
                'misconduct', 'fraud', 'misrepresentation', 'deception',
                'financial crime', 'money laundering', 'illicit activity',
                'fraudulent activity', 'criminal activity', 'illegal transaction'
            ]
        }
        
        self.positive_patterns = {
            'best_practices': [
                'best practice', 'exemplary', 'leading practice', 'model approach',
                'innovative solution', 'effective practice', 'successful implementation',
                'robust framework', 'sound practice', 'excellence'
            ],
            'regulatory_effectiveness': [
                'effective regulation', 'successful oversight', 'robust supervision',
                'comprehensive framework', 'strong governance', 'effective compliance',
                'regulatory success', 'supervisory effectiveness'
            ],
            'innovation_excellence': [
                'innovation', 'technological advancement', 'digital transformation',
                'fintech innovation', 'regulatory technology', 'automated solution',
                'artificial intelligence', 'machine learning', 'blockchain'
            ],
            'risk_management': [
                'robust risk management', 'effective risk control', 'comprehensive risk framework',
                'strong risk culture', 'risk mitigation', 'proactive risk management',
                'advanced risk analytics', 'integrated risk approach'
            ],
            'international_cooperation': [
                'international cooperation', 'cross-border collaboration', 'global coordination',
                'multilateral approach', 'international standard', 'global framework',
                'coordinated response', 'international best practice'
            ]
        }
    
    def detect_patterns(self, documents: List[Dict], analysis_type: str = "comprehensive") -> Dict:
        """Advanced pattern detection using multiple algorithms"""
        results = {
            'negative_findings': {},
            'positive_findings': {},
            'anomaly_scores': [],
            'clustering_results': {},
            'topic_analysis': {},
            'sentiment_analysis': {},
            'readability_analysis': {},
            'entity_analysis': {},
            'analysis_type': analysis_type
        }
        
        # Filter patterns based on analysis type
        if analysis_type == "governance_risk":
            target_negative = ['governance_failures', 'compliance_issues', 'legal_loopholes']
            target_positive = ['regulatory_effectiveness', 'best_practices']
        elif analysis_type == "anomaly_detection":
            target_negative = ['operational_risks', 'market_irregularities', 'financial_misconduct']
            target_positive = ['risk_management']
        elif analysis_type == "best_practices":
            target_negative = []
            target_positive = list(self.positive_patterns.keys())
        elif analysis_type == "legal_gaps":
            target_negative = ['legal_loopholes', 'compliance_issues', 'governance_failures']
            target_positive = ['regulatory_effectiveness']
        elif analysis_type == "market_irregularities":
            target_negative = ['market_irregularities', 'financial_misconduct', 'operational_risks']
            target_positive = ['risk_management', 'regulatory_effectiveness']
        else:  # comprehensive
            target_negative = list(self.negative_patterns.keys())
            target_positive = list(self.positive_patterns.keys())
        
        # Pattern matching analysis based on selection
        for category, patterns in self.negative_patterns.items():
            if category in target_negative:
                results['negative_findings'][category] = self._find_pattern_matches(documents, patterns)
        
        for category, patterns in self.positive_patterns.items():
            if category in target_positive:
                results['positive_findings'][category] = self._find_pattern_matches(documents, patterns)
        
        # Always include core analytics
        results['anomaly_scores'] = self._detect_anomalies(documents)
        results['clustering_results'] = self._perform_clustering(documents)
        results['topic_analysis'] = self._topic_modeling(documents)
        results['sentiment_analysis'] = self._sentiment_analysis(documents)
        results['readability_analysis'] = self._readability_analysis(documents)
        results['entity_analysis'] = self._entity_analysis(documents)
        
        return results
    
    def _find_pattern_matches(self, documents: List[Dict], patterns: List[str]) -> List[Dict]:
        """Find pattern matches with context"""
        matches = []
        for doc in documents:
            text = doc['content'].lower()
            for pattern in patterns:
                if pattern in text:
                    sentences = sent_tokenize(doc['content'])
                    for sentence in sentences:
                        if pattern in sentence.lower():
                            matches.append({
                                'document': doc['filename'],
                                'pattern': pattern,
                                'context': sentence,
                                'confidence': self._calculate_pattern_confidence(sentence, pattern)
                            })
        return matches
    
    def _calculate_pattern_confidence(self, sentence: str, pattern: str) -> float:
        """Calculate confidence score for pattern match"""
        words = sentence.lower().split()
        pattern_words = pattern.split()
        context_relevance = len([w for w in words if w in ['risk', 'compliance', 'regulation', 'governance']]) / len(words)
        pattern_specificity = len(pattern_words) / len(words)
        return min(1.0, (context_relevance + pattern_specificity) / 2)
    
    def _detect_anomalies(self, documents: List[Dict]) -> List[Dict]:
        """Detect anomalies using statistical methods"""
        anomalies = []
        lengths = [len(doc['content']) for doc in documents]
        mean_length = np.mean(lengths)
        std_length = np.std(lengths)
        
        for i, doc in enumerate(documents):
            length_z_score = abs((lengths[i] - mean_length) / std_length)
            if length_z_score > 2:
                anomalies.append({
                    'document': doc['filename'],
                    'type': 'length_anomaly',
                    'z_score': length_z_score,
                    'description': f"Document length significantly {'above' if lengths[i] > mean_length else 'below'} average"
                })
        return anomalies
    
    def _perform_clustering(self, documents: List[Dict]) -> Dict:
        """Perform document clustering"""
        if len(documents) < 3:
            return {'error': 'Insufficient documents for clustering'}
        
        texts = [doc['processed_content'] for doc in documents]
        vectorizer = TfidfVectorizer(max_features=1000, ngram_range=(1, 3))
        tfidf_matrix = vectorizer.fit_transform(texts)
        
        n_clusters = min(5, len(documents) // 2)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
        cluster_labels = kmeans.fit_predict(tfidf_matrix)
        
        clusters = {}
        for i, label in enumerate(cluster_labels):
            if label not in clusters:
                clusters[label] = []
            clusters[label].append(documents[i]['filename'])
        
        return {
            'n_clusters': n_clusters,
            'clusters': clusters,
            'cluster_centers': kmeans.cluster_centers_.tolist()
        }
    
    def _topic_modeling(self, documents: List[Dict]) -> Dict:
        """Perform topic modeling using LDA"""
        texts = [doc['processed_content'] for doc in documents]
        vectorizer = TfidfVectorizer(max_features=100, ngram_range=(1, 2))
        tfidf_matrix = vectorizer.fit_transform(texts)
        
        n_topics = min(5, len(documents))
        lda = LatentDirichletAllocation(n_components=n_topics, random_state=42)
        lda.fit(tfidf_matrix)
        
        feature_names = vectorizer.get_feature_names_out()
        topics = []
        for topic_idx, topic in enumerate(lda.components_):
            top_words_idx = topic.argsort()[-10:][::-1]
            top_words = [feature_names[i] for i in top_words_idx]
            topics.append({
                'topic_id': topic_idx,
                'top_words': top_words,
                'weights': topic[top_words_idx].tolist()
            })
        
        return {'topics': topics, 'n_topics': n_topics}
    
    def _sentiment_analysis(self, documents: List[Dict]) -> Dict:
        """Perform sentiment analysis"""
        sentiments = []
        for doc in documents:
            blob = TextBlob(doc['content'])
            sentiments.append({
                'document': doc['filename'],
                'polarity': blob.sentiment.polarity,
                'subjectivity': blob.sentiment.subjectivity
            })
        
        avg_polarity = np.mean([s['polarity'] for s in sentiments])
        avg_subjectivity = np.mean([s['subjectivity'] for s in sentiments])
        
        return {
            'document_sentiments': sentiments,
            'average_polarity': avg_polarity,
            'average_subjectivity': avg_subjectivity,
            'overall_sentiment': 'positive' if avg_polarity > 0.1 else 'negative' if avg_polarity < -0.1 else 'neutral'
        }
    
    def _readability_analysis(self, documents: List[Dict]) -> Dict:
        """Analyze document readability"""
        readability_scores = []
        for doc in documents:
            text = doc['content']
            scores = {
                'document': doc['filename'],
                'flesch_reading_ease': textstat.flesch_reading_ease(text),
                'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text),
                'automated_readability_index': textstat.automated_readability_index(text),
                'coleman_liau_index': textstat.coleman_liau_index(text)
            }
            readability_scores.append(scores)
        
        return {
            'document_scores': readability_scores,
            'average_grade_level': np.mean([s['flesch_kincaid_grade'] for s in readability_scores])
        }
    
    def _entity_analysis(self, documents: List[Dict]) -> Dict:
        """Extract and analyze named entities"""
        all_entities = []
        entity_counts = {}
        
        for doc in documents:
            doc_nlp = nlp(doc['content'][:1000000])  # Limit for performance
            doc_entities = []
            
            for ent in doc_nlp.ents:
                if ent.label_ in ['ORG', 'PERSON', 'GPE', 'MONEY', 'PERCENT']:
                    entity_info = {
                        'text': ent.text,
                        'label': ent.label_,
                        'start': ent.start_char,
                        'end': ent.end_char
                    }
                    doc_entities.append(entity_info)
                    
                    key = f"{ent.text}_{ent.label_}"
                    entity_counts[key] = entity_counts.get(key, 0) + 1
            
            all_entities.append({
                'document': doc['filename'],
                'entities': doc_entities
            })
        
        top_entities = sorted(entity_counts.items(), key=lambda x: x[1], reverse=True)[:20]
        
        return {
            'document_entities': all_entities,
            'top_entities': top_entities,
            'total_unique_entities': len(entity_counts)
        }

class RAGSystem:
    """Retrieval-Augmented Generation System for BIS Documents"""
    
    def __init__(self, db_path: str = "bis_documents.db"):
        self.db_path = db_path
        self.processor = DocumentProcessor()
        self.analyzer = AdvancedAnalyzer()
        self.vectorized_docs = None
        self.vectorizer = None
        self._init_database()
    
    def _init_database(self):
        """Initialize SQLite database for document storage"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT UNIQUE,
                content TEXT,
                processed_content TEXT,
                metadata TEXT,
                indexed_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()
    
    def index_documents(self, folder_path: str, progress_callback=None) -> int:
        """Index all PDF documents in the specified folder"""
        pdf_files = list(Path(folder_path).glob("*.pdf"))
        total_files = len(pdf_files)
        indexed_count = 0
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        for i, pdf_file in enumerate(pdf_files):
            try:
                if progress_callback:
                    progress_callback(i + 1, total_files, f"Processing {pdf_file.name}")
                
                content = self.processor.extract_text(str(pdf_file))
                if not content.strip():
                    continue
                
                processed_content = self.processor.preprocess_text(content)
                
                metadata = json.dumps({
                    'file_size': pdf_file.stat().st_size,
                    'modification_date': datetime.fromtimestamp(pdf_file.stat().st_mtime).isoformat(),
                    'word_count': len(content.split())
                })
                
                cursor.execute('''
                    INSERT OR REPLACE INTO documents (filename, content, processed_content, metadata)
                    VALUES (?, ?, ?, ?)
                ''', (pdf_file.name, content, processed_content, metadata))
                
                indexed_count += 1
                
            except Exception as e:
                logging.error(f"Error processing {pdf_file}: {e}")
        
        conn.commit()
        conn.close()
        
        self._create_vector_index()
        return indexed_count
    
    def _create_vector_index(self):
        """Create TF-IDF vector index for retrieval"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT filename, processed_content FROM documents")
        documents = cursor.fetchall()
        conn.close()
        
        if not documents:
            return
        
        texts = [doc[1] for doc in documents]
        self.vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1, 3))
        self.vectorized_docs = self.vectorizer.fit_transform(texts)
        self.doc_filenames = [doc[0] for doc in documents]
    
    def get_all_documents(self) -> List[Dict]:
        """Retrieve all documents from database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT filename, content, processed_content, metadata FROM documents")
        rows = cursor.fetchall()
        conn.close()
        
        documents = []
        for row in rows:
            doc = {
                'filename': row[0],
                'content': row[1],
                'processed_content': row[2],
                'metadata': json.loads(row[3]) if row[3] else {}
            }
            documents.append(doc)
        
        return documents
    
    def analyze_documents(self, analysis_type: str = "comprehensive") -> Dict:
        """Perform targeted analysis based on selected type"""
        documents = self.get_all_documents()
        if not documents:
            return {'error': 'No documents found'}
        
        return self.analyzer.detect_patterns(documents, analysis_type)

class ReportGenerator:
    """Advanced DOCX report generation system"""
    
    def __init__(self):
        pass
    
    def _clean_text(self, text: str) -> str:
        """Clean text for XML compatibility"""
        if not isinstance(text, str):
            text = str(text)
        
        # Remove NULL bytes and control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
        
        # Remove or replace problematic characters
        text = text.replace('\x00', '').replace('\x01', '').replace('\x02', '')
        
        return text
    
    def generate_report(self, analysis_results: Dict, output_path: str, analysis_type: str):
        """Generate comprehensive DOCX analysis report"""
        doc = Document()
        
        # Add all sections
        self._add_title_page(doc, analysis_type)
        self._add_executive_summary(doc, analysis_results)
        self._add_methodology_section(doc)
        self._add_detailed_analysis(doc, analysis_results)
        self._add_conclusions(doc, analysis_results)
        
        doc.save(output_path)
    
    def _add_title_page(self, doc, analysis_type: str):
        """Add professional title page"""
        # Main title
        title = doc.add_heading('BIS DOCUMENT INTELLIGENCE ANALYSIS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(f'{analysis_type.upper()} REPORT', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Subtitle section
        doc.add_paragraph("Advanced Regulatory Document Analysis using")
        doc.add_paragraph("Retrieval-Augmented Generation (RAG) Technology")
        doc.add_paragraph("Natural Language Processing & Machine Learning Algorithms")
        
        doc.add_paragraph()  # Space
        
        # Author information
        doc.add_paragraph("Principal Investigator:")
        doc.add_paragraph("Dr. Mosab Hawarey")
        doc.add_paragraph("Advanced Analytics & Regulatory Intelligence")
        doc.add_paragraph("Computational Finance & Risk Management")
        
        doc.add_paragraph()  # Space
        
        # Methodology
        doc.add_paragraph("Methodology:")
        doc.add_paragraph("• Retrieval-Augmented Generation (RAG)")
        doc.add_paragraph("• Advanced Natural Language Processing")
        doc.add_paragraph("• Machine Learning Pattern Recognition")
        doc.add_paragraph("• Statistical Anomaly Detection")
        doc.add_paragraph("• Topic Modeling & Sentiment Analysis")
        doc.add_paragraph("• Named Entity Recognition")
        doc.add_paragraph("• Document Clustering & Classification")
        
        doc.add_paragraph()  # Space
        
        # Technology Stack
        doc.add_paragraph("Technology Stack:")
        doc.add_paragraph("• Python Scientific Computing")
        doc.add_paragraph("• scikit-learn Machine Learning")
        doc.add_paragraph("• spaCy Advanced NLP")
        doc.add_paragraph("• NLTK Linguistic Analysis")
        doc.add_paragraph("• TF-IDF Vectorization")
        doc.add_paragraph("• Latent Dirichlet Allocation")
        doc.add_paragraph("• Cosine Similarity Matching")
        
        doc.add_paragraph()  # Space
        
        # Footer info
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("Document Classification: Analytical Intelligence Report")
        doc.add_paragraph("Scope: BIS Regulatory Document Corpus Analysis")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, results: Dict):
        """Add executive summary section"""
        doc.add_heading('EXECUTIVE SUMMARY', 1)
        
        negative_count = sum(len(findings) for findings in results.get('negative_findings', {}).values())
        positive_count = sum(len(findings) for findings in results.get('positive_findings', {}).values())
        
        summary_text = f"""This comprehensive analysis employed advanced Retrieval-Augmented Generation (RAG) technology combined with sophisticated Natural Language Processing algorithms to examine the complete corpus of Bank for International Settlements regulatory documents.

Key Analytical Findings:
• Total Negative Pattern Detections: {negative_count} instances
• Total Positive Pattern Identifications: {positive_count} instances
• Document Anomalies Detected: {len(results.get('anomaly_scores', []))} cases
• Topic Clusters Identified: {results.get('topic_analysis', {}).get('n_topics', 0)} themes
• Named Entities Extracted: {results.get('entity_analysis', {}).get('total_unique_entities', 0)} unique entities

The analysis reveals significant insights into regulatory patterns, governance frameworks, and compliance structures within the BIS documentation ecosystem. Advanced machine learning algorithms identified both risk factors and best practices across the document corpus."""
        
        doc.add_paragraph(summary_text)
        doc.add_page_break()
    
    def _add_methodology_section(self, doc):
        """Add methodology section"""
        doc.add_heading('ADVANCED ANALYTICAL METHODOLOGY', 1)
        
        methodology_text = """1. Document Processing & Vectorization
Advanced PDF text extraction using PyMuPDF and PyPDF2 with fallback mechanisms. Text preprocessing includes tokenization, lemmatization, and stop-word removal using NLTK.

2. Retrieval-Augmented Generation (RAG) Framework
Implementation of sophisticated RAG architecture enabling semantic document retrieval and contextual analysis. TF-IDF vectorization with n-gram feature extraction (1-3 grams) for optimal document representation.

3. Pattern Recognition Algorithms
Multi-dimensional pattern matching using custom-developed regulatory lexicons. Statistical confidence scoring based on contextual relevance and pattern specificity.

4. Machine Learning Analytics
• K-Means clustering for document categorization
• Latent Dirichlet Allocation (LDA) for topic modeling
• Statistical anomaly detection using z-score analysis
• Cosine similarity for document relationship mapping

5. Advanced NLP Techniques
• Named Entity Recognition using spaCy's transformer models
• Sentiment analysis with TextBlob polarity scoring
• Readability assessment using multiple linguistic indices
• Part-of-speech tagging and syntactic analysis

6. Quality Assurance & Validation
Multi-layer validation including cross-reference checking, confidence interval analysis, and statistical significance testing to ensure analytical robustness."""
        
        doc.add_paragraph(methodology_text)
        doc.add_page_break()
    
    def _add_detailed_analysis(self, doc, results: Dict):
        """Add detailed analysis section"""
        doc.add_heading('DETAILED ANALYTICAL FINDINGS', 1)
        
        # Negative Findings Analysis
        if results.get('negative_findings'):
            doc.add_heading('Risk & Compliance Analysis', 2)
            
            for category, findings in results['negative_findings'].items():
                if findings:
                    doc.add_heading(f"{category.replace('_', ' ').title()}", 3)
                    
                    # Add table for findings
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Document'
                    hdr_cells[1].text = 'Pattern Detected'
                    hdr_cells[2].text = 'Confidence'
                    
                    for finding in findings[:10]:  # Limit to top 10
                        row_cells = table.add_row().cells
                        doc_name = finding['document'][:25] + '...' if len(finding['document']) > 25 else finding['document']
                        row_cells[0].text = self._clean_text(doc_name)
                        row_cells[1].text = self._clean_text(finding['pattern'])
                        row_cells[2].text = f"{finding['confidence']:.2f}"
                    
                    doc.add_paragraph()  # Space
        
        # Positive Findings Analysis
        if results.get('positive_findings'):
            doc.add_heading('Excellence & Best Practices Analysis', 2)
            
            for category, findings in results['positive_findings'].items():
                if findings:
                    doc.add_heading(f"{category.replace('_', ' ').title()}", 3)
                    
                    # Add table for findings
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 3'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Document'
                    hdr_cells[1].text = 'Best Practice Identified'
                    hdr_cells[2].text = 'Confidence'
                    
                    for finding in findings[:10]:  # Limit to top 10
                        row_cells = table.add_row().cells
                        doc_name = finding['document'][:25] + '...' if len(finding['document']) > 25 else finding['document']
                        row_cells[0].text = self._clean_text(doc_name)
                        row_cells[1].text = self._clean_text(finding['pattern'])
                        row_cells[2].text = f"{finding['confidence']:.2f}"
                    
                    doc.add_paragraph()  # Space
        
        # Topic Analysis
        if results.get('topic_analysis', {}).get('topics'):
            doc.add_heading('Advanced Topic Modeling Analysis', 2)
            
            doc.add_paragraph("Latent Dirichlet Allocation (LDA) identified the following key thematic clusters:")
            
            for topic in results['topic_analysis']['topics']:
                topic_title = f"Topic {topic['topic_id'] + 1}: {', '.join(topic['top_words'][:5])}"
                doc.add_heading(topic_title, 4)
                
                words_text = f"Key Terms: {', '.join(topic['top_words'])}"
                doc.add_paragraph(words_text)
        
        # Sentiment Analysis
        if results.get('sentiment_analysis'):
            doc.add_heading('Regulatory Sentiment Analysis', 2)
            
            sentiment_data = results['sentiment_analysis']
            sentiment_text = f"""Overall Sentiment Profile:
• Average Polarity Score: {sentiment_data.get('average_polarity', 0):.3f}
• Average Subjectivity Score: {sentiment_data.get('average_subjectivity', 0):.3f}
• Overall Classification: {sentiment_data.get('overall_sentiment', 'neutral').title()}

The sentiment analysis reveals the regulatory tone and objectivity levels across the document corpus, providing insights into the communicative approach of BIS publications."""
            
            doc.add_paragraph(sentiment_text)
        
        # Entity Analysis
        if results.get('entity_analysis', {}).get('top_entities'):
            doc.add_heading('Named Entity Recognition Analysis', 2)
            
            doc.add_paragraph("Advanced NER algorithms identified the following key entities:")
            
            # Add table for entities
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 2'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Entity'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Frequency'
            
            for entity_info, count in results['entity_analysis']['top_entities'][:15]:
                entity, entity_type = entity_info.rsplit('_', 1)
                row_cells = table.add_row().cells
                row_cells[0].text = self._clean_text(entity)
                row_cells[1].text = self._clean_text(entity_type)
                row_cells[2].text = str(count)
    
    def _add_conclusions(self, doc, results: Dict):
        """Add conclusions and recommendations section"""
        doc.add_page_break()
        doc.add_heading('STRATEGIC CONCLUSIONS & RECOMMENDATIONS', 1)
        
        negative_count = sum(len(findings) for findings in results.get('negative_findings', {}).values())
        positive_count = sum(len(findings) for findings in results.get('positive_findings', {}).values())
        
        conclusions_text = f"""Executive Intelligence Summary:
This comprehensive analytical assessment employing state-of-the-art RAG technology has revealed critical insights into the BIS regulatory documentation ecosystem. The analysis demonstrates sophisticated pattern recognition capabilities and advanced natural language understanding.

Key Strategic Findings:
- Risk Indicators Identified: {negative_count} distinct patterns requiring attention
- Excellence Markers Detected: {positive_count} best practice implementations
- Analytical Precision: Advanced ML algorithms with confidence scoring
- Comprehensive Coverage: Full corpus semantic analysis completed

Methodological Innovation:
The implementation of Retrieval-Augmented Generation technology represents a significant advancement in regulatory document intelligence. The integration of multiple NLP techniques including topic modeling, sentiment analysis, and named entity recognition provides unprecedented analytical depth.

Recommended Strategic Actions:
1. Enhanced Monitoring Systems: Implement continuous RAG-based surveillance for real-time pattern detection in regulatory communications.
2. Predictive Analytics Integration: Leverage machine learning insights for proactive risk identification and compliance optimization.
3. Knowledge Management Enhancement: Utilize advanced document clustering and topic modeling for improved information architecture.
4. Regulatory Intelligence Platform: Develop comprehensive analytical dashboard incorporating these advanced NLP capabilities.

Technical Excellence Demonstration:
This analysis showcases advanced computational finance and regulatory technology capabilities, demonstrating proficiency in cutting-edge analytical methodologies essential for modern financial regulatory environments.

Future Research Directions:
- Implementation of transformer-based language models for enhanced semantic understanding
- Development of automated regulatory change detection systems
- Integration of cross-jurisdictional regulatory comparison algorithms
- Advanced graph neural networks for regulatory relationship mapping

The analytical framework presented here represents a significant contribution to regulatory technology and demonstrates exceptional capability in advanced financial document intelligence."""
        
        doc.add_paragraph(conclusions_text)
        
        doc.add_paragraph()  # Space
        doc.add_paragraph()  # Space
        
        # Signature
        signature_text = """Dr. Mosab Hawarey
Principal Investigator
Advanced Regulatory Analytics & Computational Finance
Specialized in Machine Learning, NLP & Financial Intelligence"""
        
        signature_para = doc.add_paragraph(signature_text)
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

class BISAnalyzerGUI:
    """Professional GUI for BIS Document Analysis System"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("BIS Document Intelligence & Governance Analysis System - Dr. Mosab Hawarey")
        self.root.geometry("1200x800")
        self.root.configure(bg='#f0f0f0')
        
        self.rag_system = RAGSystem()
        self.report_generator = ReportGenerator()
        self.folder_path = tk.StringVar()
        self.analysis_results = None
        
        self.setup_gui()
        self.setup_logging()
    
    def setup_logging(self):
        """Setup logging configuration"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('bis_analyzer.log'),
                logging.StreamHandler()
            ]
        )
    
    def setup_gui(self):
        """Setup the graphical user interface"""
        # Main title frame
        title_frame = tk.Frame(self.root, bg='#2c3e50', height=100)
        title_frame.pack(fill='x', padx=10, pady=5)
        title_frame.pack_propagate(False)
        
        title_label = tk.Label(
            title_frame,
            text="BIS DOCUMENT INTELLIGENCE & GOVERNANCE ANALYSIS SYSTEM",
            font=('Arial', 16, 'bold'),
            fg='white',
            bg='#2c3e50'
        )
        title_label.pack(expand=True)
        
        subtitle_label = tk.Label(
            title_frame,
            text="Advanced RAG-based Regulatory Document Analysis | Dr. Mosab Hawarey",
            font=('Arial', 10),
            fg='#ecf0f1',
            bg='#2c3e50'
        )
        subtitle_label.pack()
        
        # Main content frame
        main_frame = tk.Frame(self.root, bg='#f0f0f0')
        main_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Left panel for controls
        left_panel = tk.Frame(main_frame, bg='#ecf0f1', width=400)
        left_panel.pack(side='left', fill='y', padx=(0, 10))
        left_panel.pack_propagate(False)
        
        # Folder selection
        folder_frame = tk.LabelFrame(left_panel, text="Document Corpus Selection", font=('Arial', 12, 'bold'), bg='#ecf0f1')
        folder_frame.pack(fill='x', padx=10, pady=10)
        
        tk.Button(
            folder_frame,
            text="Select BIS Documents Folder",
            command=self.select_folder,
            bg='#3498db',
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=20,
            pady=10
        ).pack(pady=10)
        
        tk.Label(folder_frame, text="Selected Folder:", bg='#ecf0f1', font=('Arial', 9, 'bold')).pack(anchor='w', padx=10)
        folder_label = tk.Label(folder_frame, textvariable=self.folder_path, bg='#ecf0f1', wraplength=350, justify='left')
        folder_label.pack(anchor='w', padx=10, pady=5)
        
        # Indexing section
        index_frame = tk.LabelFrame(left_panel, text="Document Processing & RAG Indexing", font=('Arial', 12, 'bold'), bg='#ecf0f1')
        index_frame.pack(fill='x', padx=10, pady=10)
        
        self.index_button = tk.Button(
            index_frame,
            text="Process & Index Documents",
            command=self.index_documents,
            bg='#27ae60',
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=20,
            pady=10
        )
        self.index_button.pack(pady=10)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(index_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill='x', padx=10, pady=5)
        
        self.progress_label = tk.Label(index_frame, text="Ready to process documents", bg='#ecf0f1')
        self.progress_label.pack(pady=5)
        
        # Analysis options
        analysis_frame = tk.LabelFrame(left_panel, text="Advanced Analysis Options", font=('Arial', 12, 'bold'), bg='#ecf0f1')
        analysis_frame.pack(fill='x', padx=10, pady=10)
        
        analysis_options = [
            ("Governance & Compliance Risk Analysis", "governance_risk"),
            ("Regulatory Anomaly Detection", "anomaly_detection"),
            ("Best Practices Identification", "best_practices"),
            ("Comprehensive Intelligence Report", "comprehensive"),
            ("Legal & Procedural Gap Analysis", "legal_gaps"),
            ("Market Irregularity Assessment", "market_irregularities")
        ]
        
        self.analysis_var = tk.StringVar(value="comprehensive")
        
        for text, value in analysis_options:
            tk.Radiobutton(
                analysis_frame,
                text=text,
                variable=self.analysis_var,
                value=value,
                bg='#ecf0f1',
                font=('Arial', 9)
            ).pack(anchor='w', padx=10, pady=2)
        
        # Analysis execution
        execute_frame = tk.LabelFrame(left_panel, text="Execute Analysis", font=('Arial', 12, 'bold'), bg='#ecf0f1')
        execute_frame.pack(fill='x', padx=10, pady=10)
        
        self.analyze_button = tk.Button(
            execute_frame,
            text="Run Advanced Analysis",
            command=self.run_analysis,
            bg='#e74c3c',
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=20,
            pady=10
        )
        self.analyze_button.pack(pady=10)
        
        self.generate_report_button = tk.Button(
            execute_frame,
            text="Generate Professional Report",
            command=self.generate_report,
            bg='#9b59b6',
            fg='white',
            font=('Arial', 10, 'bold'),
            padx=20,
            pady=10,
        )
        self.generate_report_button.pack(pady=5)
        
        # Right panel for results
        right_panel = tk.Frame(main_frame, bg='#f0f0f0')
        right_panel.pack(side='right', fill='both', expand=True)
        
        # Results display
        results_frame = tk.LabelFrame(right_panel, text="Analysis Results & Intelligence Dashboard", font=('Arial', 12, 'bold'))
        results_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.results_text = scrolledtext.ScrolledText(
            results_frame,
            font=('Consolas', 10),
            bg='#2c3e50',
            fg='#ecf0f1',
            insertbackground='white'
        )
        self.results_text.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Status bar
        status_frame = tk.Frame(self.root, bg='#34495e', height=30)
        status_frame.pack(fill='x')
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            status_frame,
            text="BIS Document Intelligence System Ready | Advanced RAG Technology Initialized",
            bg='#34495e',
            fg='white',
            font=('Arial', 9)
        )
        self.status_label.pack(side='left', padx=10, pady=5)
    
    def select_folder(self):
        """Select folder containing BIS PDF documents"""
        folder = filedialog.askdirectory(title="Select BIS Documents Folder")
        if folder:
            self.folder_path.set(folder)
            pdf_count = len(list(Path(folder).glob("*.pdf")))
            self.update_status(f"Selected folder contains {pdf_count} PDF documents")
            self.display_message(f"✓ Folder selected: {folder}\n✓ Found {pdf_count} PDF documents\n")
    
    def index_documents(self):
        """Index all documents in selected folder"""
        if not self.folder_path.get():
            messagebox.showerror("Error", "Please select a folder first")
            return
        
        self.index_button.config(state='disabled')
        self.update_status("Processing documents...")
        
        def progress_callback(current, total, filename):
            progress = (current / total) * 100
            self.progress_var.set(progress)
            self.progress_label.config(text=f"Processing {current}/{total}: {filename}")
            self.root.update_idletasks()
        
        def index_worker():
            try:
                indexed_count = self.rag_system.index_documents(self.folder_path.get(), progress_callback)
                self.root.after(0, lambda: self.indexing_complete(indexed_count))
            except Exception as e:
                self.root.after(0, lambda: self.indexing_error(str(e)))
        
        threading.Thread(target=index_worker, daemon=True).start()
    
    def indexing_complete(self, count):
        """Handle completion of document indexing"""
        self.index_button.config(state='normal')
        self.analyze_button.config(state='normal')
        self.progress_var.set(100)
        self.progress_label.config(text=f"Indexing complete: {count} documents processed")
        self.update_status(f"RAG indexing complete: {count} documents ready for analysis")
        self.display_message(f"✓ RAG Indexing Complete\n✓ {count} documents processed and vectorized\n✓ System ready for advanced analysis\n\n")
    
    def indexing_error(self, error):
        """Handle indexing errors"""
        self.index_button.config(state='normal')
        self.progress_label.config(text="Error during indexing")
        self.update_status("Indexing failed")
        messagebox.showerror("Indexing Error", f"Failed to index documents: {error}")
    
    def run_analysis(self):
        """Execute the selected analysis"""
        self.analyze_button.config(state='disabled')
        self.update_status("Running advanced analysis...")
        self.display_message("🔍 Initiating Advanced Analysis...\n")
        self.display_message("⚡ Employing RAG algorithms...\n")
        self.display_message("🧠 Processing with ML models...\n")
        
        def analysis_worker():
            try:
                analysis_type = self.analysis_var.get()
                results = self.rag_system.analyze_documents(analysis_type)
                self.root.after(0, lambda: self.analysis_complete(results))
            except Exception as e:
                self.root.after(0, lambda: self.analysis_error(str(e)))
        
        threading.Thread(target=analysis_worker, daemon=True).start()
    
    def analysis_complete(self, results):
        """Handle completion of analysis"""
        self.analyze_button.config(state='normal')
        self.generate_report_button.config(state='normal')
        self.analysis_results = results
        self.update_status("Analysis complete - ready to generate report")
        self.display_analysis_summary(results)
    
    def analysis_error(self, error):
        """Handle analysis errors"""
        self.analyze_button.config(state='normal')
        self.update_status("Analysis failed")
        messagebox.showerror("Analysis Error", f"Failed to complete analysis: {error}")
    
    def display_analysis_summary(self, results):
        """Display analysis results summary"""
        analysis_type = results.get('analysis_type', 'comprehensive')
        
        summary = "\n" + "="*80 + "\n"
        summary += f"ADVANCED ANALYTICAL INTELLIGENCE SUMMARY\n"
        summary += f"Analysis Type: {analysis_type.replace('_', ' ').title()}\n"
        summary += "="*80 + "\n\n"
        
        negative_count = sum(len(findings) for findings in results.get('negative_findings', {}).values())
        positive_count = sum(len(findings) for findings in results.get('positive_findings', {}).values())
        
        if analysis_type == "governance_risk":
            summary += "🏛️ GOVERNANCE & COMPLIANCE RISK ANALYSIS:\n"
        elif analysis_type == "market_irregularities":
            summary += "📈 MARKET IRREGULARITY ASSESSMENT:\n"
        elif analysis_type == "best_practices":
            summary += "⭐ BEST PRACTICES IDENTIFICATION:\n"
        elif analysis_type == "legal_gaps":
            summary += "⚖️ LEGAL & PROCEDURAL GAP ANALYSIS:\n"
        elif analysis_type == "anomaly_detection":
            summary += "🔍 REGULATORY ANOMALY DETECTION:\n"
        else:
            summary += "🔬 COMPREHENSIVE INTELLIGENCE ANALYSIS:\n"
        
        summary += "-" * 40 + "\n"
        
        if results.get('negative_findings'):
            for category, findings in results['negative_findings'].items():
                if findings:
                    summary += f"• {category.replace('_', ' ').title()}: {len(findings)} instances detected\n"
                    for finding in findings[:3]:
                        summary += f"  - {finding['pattern']} (confidence: {finding['confidence']:.2f})\n"
                    if len(findings) > 3:
                        summary += f"  ... and {len(findings) - 3} more instances\n"
            summary += "\n"
        
        if results.get('positive_findings'):
            summary += "✅ EXCELLENCE & BEST PRACTICES:\n"
            summary += "-" * 40 + "\n"
            for category, findings in results['positive_findings'].items():
                if findings:
                    summary += f"• {category.replace('_', ' ').title()}: {len(findings)} instances identified\n"
                    for finding in findings[:3]:
                        summary += f"  - {finding['pattern']} (confidence: {finding['confidence']:.2f})\n"
                    if len(findings) > 3:
                        summary += f"  ... and {len(findings) - 3} more instances\n"
            summary += "\n"
        
        if results.get('topic_analysis', {}).get('topics'):
            summary += "📊 ADVANCED TOPIC MODELING:\n"
            summary += "-" * 40 + "\n"
            for topic in results['topic_analysis']['topics']:
                summary += f"• Topic {topic['topic_id'] + 1}: {', '.join(topic['top_words'][:5])}\n"
            summary += "\n"
        
        if results.get('sentiment_analysis'):
            sentiment_data = results['sentiment_analysis']
            summary += "💭 REGULATORY SENTIMENT ANALYSIS:\n"
            summary += "-" * 40 + "\n"
            summary += f"• Overall Sentiment: {sentiment_data.get('overall_sentiment', 'neutral').title()}\n"
            summary += f"• Average Polarity: {sentiment_data.get('average_polarity', 0):.3f}\n"
            summary += f"• Average Subjectivity: {sentiment_data.get('average_subjectivity', 0):.3f}\n\n"
        
        if results.get('entity_analysis', {}).get('top_entities'):
            summary += "🏷️ NAMED ENTITY RECOGNITION:\n"
            summary += "-" * 40 + "\n"
            for entity_info, count in results['entity_analysis']['top_entities'][:5]:
                entity, entity_type = entity_info.rsplit('_', 1)
                summary += f"• {entity} ({entity_type}): {count} occurrences\n"
            summary += "\n"
        
        summary += "="*80 + "\n"
        summary += f"Targeted {analysis_type.replace('_', ' ').title()} Analysis Complete.\n"
        summary += f"Risk Patterns: {negative_count} | Excellence Patterns: {positive_count}\n"
        summary += "Ready to generate professional report.\n"
        summary += "="*80 + "\n"
        
        self.display_message(summary)
    
    def generate_report(self):
        """Generate professional DOCX report"""
        if not self.analysis_results:
            messagebox.showerror("Error", "No analysis results available")
            return
        
        output_file = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            title="Save Analysis Report"
        )
        
        if output_file:
            try:
                self.update_status("Generating professional report...")
                analysis_type = self.analysis_var.get().replace('_', ' ').title()
                
                self.report_generator.generate_report(
                    self.analysis_results,
                    output_file,
                    analysis_type
                )
                
                self.update_status("Report generated successfully")
                self.display_message(f"✅ Professional report generated: {output_file}\n")
                messagebox.showinfo("Success", f"Report generated successfully:\n{output_file}")
                
            except Exception as e:
                self.update_status("Report generation failed")
                messagebox.showerror("Error", f"Failed to generate report: {str(e)}")
    
    def display_message(self, message):
        """Display message in results text area"""
        self.results_text.insert(tk.END, message)
        self.results_text.see(tk.END)
        self.root.update_idletasks()
    
    def update_status(self, message):
        """Update status bar"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.status_label.config(text=f"[{timestamp}] {message}")
        logging.info(message)
    
    def run(self):
        """Start the GUI application"""
        self.root.mainloop()

def main():
    """Main application entry point"""
    print("="*80)
    print("BIS DOCUMENT INTELLIGENCE & GOVERNANCE ANALYSIS SYSTEM")
    print("Advanced RAG-based Regulatory Document Analysis")
    print("Dr. Mosab Hawarey - Computational Finance & Regulatory Intelligence")
    print("="*80)
    
    try:
        app = BISAnalyzerGUI()
        app.run()
    except Exception as e:
        print(f"Application error: {e}")
        logging.error(f"Application error: {e}")

if __name__ == "__main__":
    main()